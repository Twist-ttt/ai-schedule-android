# -*- coding: utf-8 -*-
"""把 docs/*.md 转成符合作业要求的 .docx（中文友好：宋体正文/微软雅黑标题、表格、页码页眉）。
用法：python md2docx.py
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_BODY = '宋体'           # 正文东亚字体
FONT_HEAD = '微软雅黑'        # 标题东亚字体
FONT_ASCII = 'Times New Roman'  # 正文西文
MONO = 'Consolas'            # 行内/块代码

INLINE_RE = re.compile(r'(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\))')


def set_run(run, size=12, bold=False, color=None, east=FONT_BODY, ascii_font=FONT_ASCII):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)


def add_inline(paragraph, text, size=12, bold=False):
    """解析 **加粗** `代码` [文字](链接) 三种行内格式。"""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            set_run(paragraph.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            set_run(paragraph.add_run(part[1:-1]), size=size - 0.5, east=FONT_BODY, ascii_font=MONO)
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[(.+?)\]\((.+?)\)', part)
            set_run(paragraph.add_run(m.group(1) if m else part), size=size, color=(0x1A, 0x56, 0xDB))
        else:
            set_run(paragraph.add_run(part), size=size, bold=bold)


def set_left_border(paragraph, color='6C63FF', sz='18'):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pbdr.append(left)
    pPr.append(pbdr)


def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('第 ')
    set_run(run, size=9, color=(0x80, 0x80, 0x80))
    # PAGE 域
    for t, is_field in (('PAGE', True),):
        f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
        r = p.add_run(); r._r.append(f1); r._r.append(it); r._r.append(f2)
        set_run(r, size=9, color=(0x80, 0x80, 0x80))
    set_run(p.add_run(' 页'), size=9, color=(0x80, 0x80, 0x80))


def add_header(section, title):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run(title), size=9, color=(0x99, 0x99, 0x99))


def convert(md_path, docx_path, doc_title):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().replace('\r\n', '\n').split('\n')

    doc = Document()
    # Normal 默认字体
    normal = doc.styles['Normal']
    normal.font.name = FONT_ASCII
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), FONT_BODY)

    # A4 + 页边距
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(25)
    sec.left_margin = sec.right_margin = Mm(28)
    add_page_number(sec)
    add_header(sec, doc_title)

    sizes = {1: 20, 2: 16, 3: 14, 4: 13, 5: 12, 6: 12}
    colors = {1: (0x1F, 0x3A, 0x5F), 2: (0x2B, 0x4C, 0x7E), 3: (0x3A, 0x5F, 0x9A)}

    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()

        # 代码块
        if s.startswith('```'):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            set_run(p.add_run('\n'.join(buf)), size=10, ascii_font=MONO, east=FONT_BODY,
                    color=(0x33, 0x33, 0x33))
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)', s)
        if m:
            lvl, text = len(m.group(1)), m.group(2)
            try:
                p = doc.add_paragraph(style=f'Heading {lvl}')  # 保留大纲级别，可生成目录
            except KeyError:
                p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14 if lvl <= 2 else 8)
            p.paragraph_format.space_after = Pt(6)
            set_run(p.add_run(text), size=sizes[lvl], bold=True,
                    color=colors.get(lvl), east=FONT_HEAD, ascii_font=FONT_HEAD)
            i += 1; continue

        # 水平线
        if s in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
            pbdr.append(bottom); pPr.append(pbdr)
            i += 1; continue

        # 表格（表头行 + 紧跟分隔行 |---|）
        if s.startswith('|') and i + 1 < n and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
            tbl = []
            while i < n and lines[i].strip().startswith('|'):
                tbl.append(lines[i].strip()); i += 1
            header = [c.strip() for c in tbl[0].strip('|').split('|')]
            data = [[c.strip() for c in r.strip('|').split('|')] for r in tbl[2:]]
            ncol = len(header)
            table = doc.add_table(rows=1, cols=ncol)
            table.style = 'Table Grid'
            table.autofit = True
            for j, h in enumerate(header):
                shade = OxmlElement('w:shd'); shade.set(qn('w:fill'), 'E8EEF7')
                cell = table.rows[0].cells[j]
                cell._tc.get_or_add_tcPr().append(shade)
                add_inline(cell.paragraphs[0], h, size=10.5, bold=True)
            for row in data:
                cells = table.add_row().cells
                for j in range(min(ncol, len(row))):
                    add_inline(cells[j].paragraphs[0], row[j], size=10.5)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # 引用块（连续 > 行合并）
        if s.startswith('>'):
            buf = []
            while i < n and lines[i].strip().startswith('>'):
                t = lines[i].strip().lstrip('>').strip()
                if t:
                    buf.append(t)
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            set_left_border(p)
            add_inline(p, ' '.join(buf) if len(buf) == 1 else buf[0], size=11)
            for extra in buf[1:]:
                ep = doc.add_paragraph()
                ep.paragraph_format.left_indent = Pt(14)
                set_left_border(ep)
                add_inline(ep, extra, size=11)
            continue

        # 无序列表
        m = re.match(r'^[-*]\s+(.*)', s)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, m.group(1), size=12)
            i += 1; continue

        # 有序列表
        m = re.match(r'^\d+\.\s+(.*)', s)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, m.group(1), size=12)
            i += 1; continue

        # 空行
        if s == '':
            i += 1; continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.4
        add_inline(p, s, size=12)
        i += 1

    doc.save(docx_path)
    print('OK ->', docx_path, '| 段落:', len(doc.paragraphs), '| 表格:', len(doc.tables))


if __name__ == '__main__':
    convert('docs/设计文档.md', '项目设计文档.docx', 'AI日程转换助手 · 项目设计文档')
    convert('docs/prompts.md', 'AI提示词使用记录.docx', 'AI日程转换助手 · AI 提示词使用记录')
